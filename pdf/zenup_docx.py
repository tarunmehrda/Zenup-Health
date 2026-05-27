import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = create_element('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def generate_document():
    doc = docx.Document()
    
    # Page setup - Standard 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Configure base font styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x37, 0x41, 0x51) # Dark grey

    # Configure custom Heading 1
    style_h1 = doc.styles['Heading 1']
    h1_font = style_h1.font
    h1_font.name = 'Calibri'
    h1_font.size = Pt(18)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(0x1A, 0x2E, 0x28) # Slate Green/Dark color
    
    # Configure custom Heading 2
    style_h2 = doc.styles['Heading 2']
    h2_font = style_h2.font
    h2_font.name = 'Calibri'
    h2_font.size = Pt(14)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(0x1D, 0xB8, 0x7E) # Accent Green
    
    # Configure custom Heading 3
    style_h3 = doc.styles['Heading 3']
    h3_font = style_h3.font
    h3_font.name = 'Calibri'
    h3_font.size = Pt(12)
    h3_font.bold = True
    h3_font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    # 1. COVER PAGE (Simple layout, no colorful backgrounds/banners)
    # Spacer at top
    for _ in range(3):
        doc.add_paragraph()
        
    p_main_title = doc.add_paragraph()
    p_main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_main_title = p_main_title.add_run("Zenup Health")
    run_main_title.font.size = Pt(36)
    run_main_title.font.bold = True
    run_main_title.font.color.rgb = RGBColor(0x1A, 0x2E, 0x28)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("75-Day Flutter App Development Plan")
    run_sub.font.size = Pt(16)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(0x7A, 0x9E, 0x93)

    p_stack = doc.add_paragraph()
    p_stack.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_stack = p_stack.add_run("Frontend Client · Flutter 3.x · iOS & Android")
    run_stack.font.size = Pt(11)
    run_stack.font.italic = True

    for _ in range(6):
        doc.add_paragraph()

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = p_meta.add_run(
        "Psychiatry · Therapy · Mental Health Booking\n"
        "Built for zenuphealth.com\n"
        "Project Manager: Samarth Mehta\n"
        "Developers: Arnav, Tarun, Yashwanth"
    )
    run_meta.font.size = Pt(10)
    run_meta.font.color.rgb = RGBColor(0x7A, 0x9E, 0x93)

    doc.add_page_break()

    # 2. TABLE OF CONTENTS (Clean textual layout)
    doc.add_heading("Table of Contents", level=1)
    toc_p = doc.add_paragraph()
    toc_items = [
        ("Project Overview & Frontend Architecture", "Page 3"),
        ("Phase 1 — Days 1–9: Foundation & Setup", "Page 4"),
        ("Phase 2 — Days 10–26: Core Features & Booking Flow", "Page 6"),
        ("Phase 3 — Days 27–45: Advanced Telehealth & Local Persistence", "Page 10"),
        ("Phase 4 — Days 46–62: Polish & Testing", "Page 15"),
        ("Phase 5 — Days 63–75: Launch & Deploy", "Page 20"),
        ("Daily Deliverables Tracker", "Page 23"),
        ("Estimated Project Costs", "Page 25"),
        ("Risk Register & Launch Checklist", "Page 26"),
    ]
    for title, pg in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6))
        p.add_run(title)
        p.add_run(f"\t{pg}").bold = True
        
    doc.add_page_break()

    # 3. OVERVIEW
    doc.add_heading("Project Overview", level=1)
    doc.add_paragraph(
        "Zenup Health is a mental health platform offering psychiatry, therapy, "
        "and counseling booking at zenuphealth.com. This Flutter app is developed as a "
        "frontend-first client utilizing client-side state management (Riverpod), local data caching "
        "(Hive), and a  Supabase backend for user authentication and core data persistence. "
        "The architecture keeps the app fast and offline-capable, while Supabase handles secure login "
        "and syncing confirmed bookings and user profiles to the cloud."
    )
    
    doc.add_heading("Project Team & Personnel", level=3)
    team_table = doc.add_table(rows=5, cols=2)
    team_table.style = 'Table Grid'
    team_hdr = team_table.rows[0].cells
    team_hdr[0].text = "Role"
    team_hdr[1].text = "Member Name"
    for cell in team_hdr:
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_background(cell, "E8F7F0")
        
    team_rows = [
        ["Head of Engineering / PM", "Samarth Mehta"],
        ["App Developer", "Arnav"],
        ["App Developer", "Tarun"],
        ["App Developer", "Yashwanth"],
    ]
    for idx, row in enumerate(team_rows):
        cells = team_table.rows[idx+1].cells
        cells[0].text = row[0]
        cells[1].text = row[1]

    doc.add_paragraph()

    doc.add_heading("Architecture Principles", level=3)
    why_items = [
        "Frontend-First: Full UI built and validated offline using mock repositories before backend wiring.",
        "Local Caching: Hive database for encrypted on-device data persistence — appointments, profiles, messages.",
        "Supabase Auth ( Backend): Supabase handles secure user sign-up, login, and session tokens only.",
        "Supabase Data Sync: Confirmed bookings and user profiles are written to Supabase Postgres in real time.",
        "Robust State Management: Riverpod providers cleanly separate UI from data layers.",
        "Offline-Capable: All screens function with local Hive cache — Supabase syncs in the background."
    ]
    for item in why_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    doc.add_heading("Complete Tech Stack", level=3)
    tech_table = doc.add_table(rows=1, cols=3)
    tech_table.style = 'Table Grid'
    tech_hdr = tech_table.rows[0].cells
    tech_hdr[0].text = "Layer"
    tech_hdr[1].text = "Technology"
    tech_hdr[2].text = "Notes"
    for cell in tech_hdr:
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_background(cell, "E8F7F0")

    tech_raw = [
        ['Frontend (App)', 'Flutter 3.x (Dart)', 'Single codebase — iOS & Android'],
        ['Backend', 'Supabase (Auth + Postgres)', 'User login, session tokens & confirmed booking sync'],
        ['State Mgmt', 'Riverpod 2.x + AsyncNotifier', 'Reactive, testable, mock-compatible state'],
        ['Local Storage', 'Hive (Encrypted)', 'Encrypted offline cache & user preferences'],
        ['Navigation', 'GoRouter', 'Declarative routing + deep links'],
        ['Payments', 'Razorpay SDK (Client)', 'Client-side checkout integration sandbox'],
        ['Video Call', 'MediaSFU SDK (Client)', 'Telehealth video & audio calling interface'],
        ['Notifications', 'Local Notifications SDK', 'Scheduled check-in & app reminder alerts'],
        ['Mocking Layer', 'Mock Repository Adapters', 'Simulated network latencies and operation buffers'],
        ['Testing', 'Flutter Test + Patrol', 'Unit, widget, golden & integration E2E tests']
    ]
    for row in tech_raw:
        row_cells = tech_table.add_row().cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1]
        row_cells[2].text = row[2]
        
    doc.add_page_break()

    # Define all phases data
    p1_data = [
        ("Day 1", "Project Init & Folder Architecture",
         ["Create Flutter project using standard naming and organization settings",
          "Establish clean architecture directory structure (features, core, shared, mock_services)",
          "Initialize Git repository and configure Git ignore files for sensitive environments",
          "Add package dependencies including Riverpod, Hive local storage, and GoRouter",
          "Configure Android and iOS project settings for minimum SDK versions"],
         ["Flutter project boots on emulator", "Git repository initialized", "Dependencies configuration complete"]),
        ("Day 2", "Local Storage & State Management Setup",
         ["Configure Riverpod global state container provider wrapper",
          "Set up application config file for environment switches (mock/dev)",
          "Initialize Hive local storage in the main application entry point",
          "Create a dedicated local storage service class wrapper using Hive",
          "Verify local cache read and write flows using test key-value pairs"],
         ["Local storage service initialized", "Riverpod state notifier configured", "Cache read/write test passing"]),
        ("Day 3", "Local Database Schema & Hive Adapters",
         ["Define Dart models representing User, Appointment, Mood, and Chat data structures",
          "Register Hive type adapters and generate Hive adapters using build_runner",
          "Create client-side data schema definitions for local models",
          "Implement local database manager with secure AES encryption keys",
          "Configure local data backup mechanisms and clean-up routines"],
         ["Hive adapters generated successfully", "Local models registered", "Local data security layer active"]),
        ("Day 4", "Design System & Theme",
         ["Define color, typography, and spacing constants in theme config",
          "Configure Material theme with selected typography and styling guidelines",
          "Build reusable custom UI components: buttons, cards, inputs, badges, and avatars",
          "Set up application asset directories for assets and icons",
          "Add branding icons and image assets to project folders"],
         ["Theme configuration complete", "Reusable UI widgets built", "Fonts rendering correctly"]),
        ("Day 5", "Navigation Configuration",
         ["Set up declarative routing using GoRouter package",
          "Define application routing paths (splash, login, dashboard, profile, booking)",
          "Implement navigation guard checking user client-side login state",
          "Create bottom navigation bar layout linking main tabs",
          "Configure deep linking pathways for application redirection"],
         ["All routes defined", "Client-side login guard redirecting correctly", "Bottom nav working"]),
        ("Day 6", "Authentication — Supabase Auth Integration",
         ["Initialize Supabase Flutter SDK and configure project URL + anon key",
          "Build login and sign-up UI with email/password validation forms",
          "Connect login form to Supabase Auth — supabase.auth.signInWithPassword()",
          "Connect sign-up form to Supabase Auth — supabase.auth.signUp()",
          "Build mock Google OAuth sign-in flow for visual feedback",
          "Handle Supabase auth error responses with user-friendly messages"],
         ["Supabase Auth login functional", "Sign-up flow creating real user in Supabase", "Auth error handling active"]),
        ("Day 7", "Auth State & Profile Sync to Supabase",
         ["Listen to Supabase auth state stream and auto-route user on login/logout",
          "On first login: show profile setup form (name, date of birth, contact)",
          "Save completed profile to Supabase 'profiles' table via REST insert",
          "Cache profile data locally in encrypted Hive box for offline access",
          "Implement local file picker and upload avatar to Supabase Storage bucket"],
         ["Supabase auth stream auto-navigating", "Profile saved to Supabase profiles table", "Avatar uploaded to Supabase Storage"]),
        ("Day 8", "Splash, Onboarding & Local State Tests",
         ["Build animated splash screen with entrance animations",
          "Create multi-slide onboarding flow with local state storage",
          "Write unit tests for authentication notifier and local storage service",
          "Fix Phase 1 visual bugs and layout adjustments",
          "Document environment settings and local storage structure in README"],
         ["Splash + onboarding working", "Unit tests passing", "README updated with local setup instructions"]),
        ("Day 9", "Buffer Day — Setup & Environment Contingency",
         ["Resolve unexpected issues with local storage initialization or path provider",
          "Verify Hive model migrations and key security setup on test platforms",
          "Debug dependency conflicts or compiler issues on target platform simulators",
          "Update local development environment onboarding documentation"],
         ["Verified local environment with updated configuration", "Setup and storage troubleshooting notes completed"]),
    ]

    p2_data = [
        ("Day 10", "Home Dashboard",
         ["Build Home screen with user greeting loading profile data",
          "Create upcoming appointments widget querying local logs in Hive",
          "Add quick-action pathways for booking, messages, and appointments overview",
          "Create featured providers listing sorted by highest ratings",
          "Add appointment scheduling summary widget using local sources"],
         ["Home screen loading cached user data", "Appointments widget showing", "Featured providers list complete"]),
        ("Day 11", "Provider Discovery & Search",
         ["Create provider directory listing filtering by specialty",
          "Implement client-side search functionality over provider names and bios",
          "Design filter selection chips for medical categories",
          "Design provider details card using profile and image resources",
          "Implement sorting options by rating and availability slots"],
         ["Provider list loaded from local assets", "Search and filtering functional", "Sorting options active"]),
        ("Day 12", "Provider Profile Screen",
         ["Fetch detailed provider profile data and reviews from local JSON assets",
          "Display bio details, credentials, ratings, and feedback counts",
          "Create availability calendar displaying open booking slots",
          "Render reviews list ordered by submission date",
          "Add booking action button redirecting to scheduling flow"],
         ["Provider details rendering correctly", "Availability slots loaded", "Reviews feed functional"]),
        ("Day 13", "Booking Flow Steps 1 & 2",
         ["Implement multi-step booking state notifier using Riverpod",
          "Step 1: Select preferred appointment medium (video, in-person, audio)",
          "Step 2: Display date picker and fetch corresponding open slots",
          "Create interactive time-slot grid disabling filled appointments",
          "Store selected slot reference in scheduling state"],
         ["Booking state managed reactively", "Date and slot selection complete", "Occupied slots disabled"]),
        ("Day 14", "Booking Flow Steps 3 & 4 — Confirm & Sync",
         ["Step 3: Collect billing information and pre-fill from user profile",
          "Step 4: Display booking summary review screen",
          "On confirm: Write appointment record to Supabase 'appointments' table",
          "Cache confirmed appointment in local Hive box for offline access",
          "Mark the selected slot as booked in the local calendar state",
          "Trigger local booking success notification reminder"],
         ["Appointment saved to Supabase appointments table", "Cached in Hive for offline", "Local confirmation alert triggered"]),
        ("Day 15", "Appointments Management Screen",
         ["Create appointments interface with tabs for upcoming, past, and cancelled",
          "Connect dashboard with local appointments stream listener for live status updates",
          "Implement reschedule option updating local booking slot and status",
          "Implement cancel option updating booking status in local storage",
          "Add cancellation policy modal warning users before confirmation"],
         ["Appointments list updating in real time", "Reschedule request processing", "Cancellation flow active"]),
        ("Day 16", "Client-Side Payment Gateway Integration",
         ["Configure Razorpay SDK package for client checkout flow",
          "Initialize client payment parameters and trigger checkout request from Flutter",
          "Design mock payment checkout screen with secure Razorpay branding",
          "Handle payment checkout success, failure, and cancellation states in UI",
          "Update local appointment record with transaction ID and status"],
         ["Razorpay SDK client integrated", "Mock checkout screen functional", "Payment status saved in local database"]),
        ("Day 17", "Billing & Invoices UI",
         ["Manage member insurance detail fields within user profile form",
          "Retrieve paid appointments list to render billing history",
          "Implement client-side PDF document generator in Flutter",
          "Generate invoice and receipt PDFs directly on the device",
          "Add insurance information edit form with form validation rules"],
         ["Insurance fields editable in profile", "Billing history populated", "Client-side PDF invoice generation functional"]),
        ("Day 18", "Client-Side Messaging UI",
         ["Build chat interface between patient and provider",
          "Load historical message records from local Hive store",
          "Implement send message operation saving new message to Hive",
          "Simulate chat stream with automated responses for demo purposes",
          "Configure typing indicators using local timer triggers"],
         ["Message history loading from Hive", "Message sending operational", "Typing indicators simulated"]),
        ("Day 19", "Local Scheduled Notifications",
         ["Configure local notification scheduler package",
          "Schedule automated reminder triggers for upcoming appointments",
          "Set up daily check-in prompt notifications at custom times",
          "Register local notification channels for Android and iOS",
          "Set up deep links to open corresponding screens when tapping notification"],
         ["Local scheduler configured", "Automated reminders working", "Notification deep links active"]),
        ("Day 20", "Buffer Day — Payment SDK Sandbox & Theme Calibration",
         ["Troubleshoot Razorpay SDK integration or layout anomalies",
          "Test simulated slow state changes and loading spinners during checkout",
          "Adjust app color palette and dark/light theme consistency",
          "Verify local state persistence after cold restarts"],
         ["Razorpay sandbox flow verified", "Layout theme audits completed"]),
        ("Day 21", "Multi-Provider Schedule Sync",
         ["Design Dart data structures representing multi-provider slots and schedule sync",
          "Create local database model for tracking schedule changes and updates",
          "Implement logic to prevent double-booking on overlapping provider slots"],
         ["Provider sync models initialized", "Schedule overlap resolver class created"]),
        ("Day 22", "Booking Customization UI",
         ["Build custom notes field during booking to add patient concerns",
          "Implement selector for language preference or doctor gender flags",
          "Design validation rules for booking custom fields"],
         ["Booking concerns field functional", "Custom doctor flags verified"]),
        ("Day 23", "UI Concurrency Guards & Input Validation",
         ["Implement client-side transaction locks to prevent double-booking on quick double-taps",
          "Design input guard validation rules for all interactive forms",
          "Write widget tests simulating rapid simultaneous user inputs on calendar slots"],
         ["Client-side booking locking logic operational", "Double-tap form validations active"]),
        ("Day 24", "In-App Booking Cancellation & Client Refund Simulation",
         ["Build slot release animation on booking cancellation",
          "Simulate wallet credits refund on cancellations within policy limits",
          "Set up local notifications alerting user of slot release"],
         ["Local refund/credit simulation active", "Slot release state updated in Hive"]),
        ("Day 25", "Client-Side State Stress Testing",
         ["Perform memory stress testing with large local datasets (1000+ mock slots)",
          "Verify UI rendering frame rates and scroll performance",
          "Optimize local search index queries over provider profiles"],
         ["Stress testing completed", "Local search index latency minimized"]),
        ("Day 26", "Local PDF Invoice Export",
         ["Automate local patient invoice PDF generation post-booking",
          "Map billing details and invoice download paths to local appointment histories",
          "Build manual invoice download button in patient's billing screen"],
         ["Invoice PDF export functional", "Billing history download link verified"]),
    ]

    p3_data = [
        ("Day 27", "Telehealth Video UI (MediaSFU Integration)",
         ["Add video consultation SDK module to the application dependencies",
          "Configure video calling client connection details for audio and video rooms",
          "Build telehealth call screen displaying remote and local camera streams",
          "Add interface controllers for camera toggling, mic mute, and screen sharing",
          "Store mock session credentials securely within local appointment notes"],
         ["Telehealth video call working on emulator", "Connected to media room layout", "Telehealth call controls functional"]),
        ("Day 28", "Telehealth — Device Testing & Polish",
         ["Test telehealth call screen on physical iOS and Android hardware",
          "Add network connection quality indicators and status overlay to call screen",
          "Build in-app post-session feedback survey",
          "Manage application backgrounding states (Picture-in-Picture mode) during live calls",
          "Integrate interactive chat panel overlay into the active video session"],
         ["Video UI tested on real devices", "Post-session survey interface complete", "In-call chat overlay working"]),
        ("Day 29", "Prescription Tracking",
         ["Create prescriptions screen showing active user prescriptions",
          "Display detail card for medication names, dosages, schedules, and providers",
          "Implement refill request flow with mock request notifications in app",
          "Configure local system notifications to remind users to take medication",
          "Build listing filter displaying archived or inactive prescriptions"],
         ["Prescriptions list rendering", "Refill request flow functional", "Local medication reminders"]),
        ("Day 30", "Health Records File Explorer",
         ["Build personal health records file explorer interface",
          "Implement local file picker to select health documents",
          "Encrypt documents before saving to secure local application directory",
          "Retrieve listing of files registered under the user's local directory",
          "Generate secure file preview screen inside the application"],
         ["Local health records explorer complete", "Files encrypted and stored", "In-app file preview working"]),
        ("Day 31", "Advanced Booking Filters",
         ["Build search filter form allowing sorting of doctor directory by availability timeslots",
          "Create active filters status bar widget for selected criteria",
          "Build interactive filters resetting animation",
          "Implement favorites list caching doctors using Hive database box",
          "Generate provider rating summaries calculating score stats locally"],
         ["Favorites system registered in Hive", "Availability filters functional", "Dynamic search active"]),
        ("Day 32", "In-App Alerts & Settings",
         ["Build list showing local notifications and system alerts",
          "Create account settings screen allowing users to edit profile and reset password",
          "Configure notification preference switches and save choices in local profile",
          "Integrate option for analytics collection opt-out in settings",
          "Add local data purging tool triggering secure storage cleanup"],
         ["Notifications center screen active", "Settings updating preferences in Hive", "Data purging flow functional"]),
        ("Day 33", "Provider Rating Form",
         ["Build rating feedback form linking patient and provider",
          "Configure validation to only allow feedback for completed appointments",
          "Recalculate average provider rating locally when new feedback is submitted",
          "Update provider local cache database in Hive",
          "Display updated review feed on the provider details screen"],
         ["Review form UI complete", "Provider rating auto-calculated locally", "Reviews visible on details profile"]),
        ("Day 34", "Local Appointment Alerts & Campaign Prompts",
         ["Create scheduled timer checking for upcoming 24-hour bookings",
          "Trigger local notification alert for upcoming patient slots",
          "Implement in-app campaign banner prompts suggesting scheduling",
          "Verify user notification preferences before showing campaign banners"],
         ["Timer check active", "24-hour alerts functional", "Campaign prompts working"]),
        ("Day 35", "Medication Reminder Alerts & Logs",
         ["Build scheduled timers for medication schedule intakes",
          "Create localized daily logs tracking consumed doses",
          "Verify medication alerts triggers using background task scheduler"],
         ["Medication reminders working", "Push alarms configured", "Intake logs saved locally"]),
        ("Day 36", "Telehealth Connection Diagnostics UI",
         ["Design call diagnostic overlay alerting user of network issues",
          "Log call exception events in local diagnostic logs",
          "Set up automatic rescheduling recommendation dialogs for patients"],
         ["Diagnostics overlay functional", "Call failure logging active", "Reschedule dialog active"]),
        ("Day 37", "Post-Call Consultation Summary & Review",
         ["Trigger post-consultation feedback prompt on call end",
          "Sync review score and comments to local appointment history",
          "Design escalation dialog for negative rating feedback"],
         ["Post-call prompt active", "Reviews logged in local history", "Escalation dialog verified"]),
        ("Day 38", "Offline Outbox Queue & Sync",
         ["Implement offline sync manager caching user actions while offline",
          "Display connectivity banner when app is offline",
          "Implement retry logic to process outbox on reconnect",
          "Clear outbox queue database on logout"],
         ["Outbox queue active", "Offline banner working", "Cache cleared on logout"]),
        ("Day 39", "Referral Sharing UI",
         ["Add referral code generation logic based on user profile metadata",
          "Generate unique alphanumeric code for each user on sign-up",
          "Enable referral code sharing via native share utility (share_plus)",
          "Apply registration attribution logic to credit referred users",
          "Create rewards display showing total counts of successful referrals"],
         ["Referral code generated", "Native share sheet working", "Rewards dashboard complete"]),
        ("Day 40", "Accessibility & Localization",
         ["Audit application layout against accessibility standard checklists",
          "Add semantic labels and accessibility tags to all interactive components",
          "Configure localization libraries for multi-language display support",
          "Create translated text resource files for supported languages",
          "Test application interface using system screen reader tools"],
         ["Accessibility audit complete", "Screen reader working", "Spanish translations added"]),
        ("Day 41", "Security Hardening & Encryption",
         ["Verify that all local database boxes in Hive use AES-256 encryption keys",
          "Implement local biometric authentication check (fingerprint/face)",
          "Encrypt local file directories containing personal health documents",
          "Add secure SSL pinning configuration for external API calls",
          "Audit local security parameters in build configurations"],
         ["Local database encrypted", "Biometric auth gate working", "Folder encryption active"]),
        ("Day 42", "Phase 3 Integration Test & Demo",
         ["Run full client scenario flow testing all advanced features",
          "Profile memory performance and resolve any identified leaks",
          "Review local query execution latency and optimize list indexing",
          "Conduct system walkthrough demo highlighting completed functionality"],
         ["Integration tests passing", "No memory leaks", "Walkthrough demo verified"]),
        ("Day 43", "Buffer Day — Telehealth Call Screens & Layout Adjustments",
         ["Debug device hardware permission flows (mic/camera) on physical target platforms",
          "Verify offline storage caching and outbox sync queue on flaky networks",
          "Resolve layout responsiveness issues for telehealth overlay screens on tablets"],
         ["Offline outbox verified", "Telehealth call overlays tested on tablet layouts"]),
        ("Day 44", "Local Cache Pruning & Purging Utility",
         ["Implement memory cleanup routines deleting expired schedules and logs",
          "Create diagnostic dashboard showing cache size",
          "Verify that GDPR-compliant profile and appointment purge flows delete completely"],
         ["Local cache statistics operational", "Data cache purging utility active", "GDPR deletion verified"]),
        ("Day 45", "Offline Flow Simulation Dry-Run",
         ["Simulate extended offline session with multi-booking and messaging flows",
          "Audit local outbox queue storage limits and retry counts",
          "Conduct offline-to-online reconciliation dry-run test"],
         ["Offline flow simulations complete", "Offline sync outbox validated", "Reconciliation logic verified"]),
    ]

    p4_data = [
        ("Day 46", "UI Animation Polish",
         ["Hero transitions between provider list and provider profile screens",
          "Shimmer loading placeholders on all list screens during data fetch",
          "Lottie animations for empty states (no appointments, no messages)",
          "Micro-interactions: button press haptic feedback, card tap scale",
          "Smooth PageView transitions in onboarding flows"],
         ["Hero transitions working", "Shimmer on all loaders active", "Lottie empty states rendering"]),
        ("Day 47", "Error Handling & Empty States",
         ["Wrap all repository calls in try/catch blocks with custom exceptions",
          "Show user-friendly, localized error messages in UI",
          "Empty state layouts for: no appointments, no messages, no prescriptions",
          "Full-screen error widget with retry button that re-calls service notifier",
          "Network connectivity lost banner via connectivity_plus package"],
         ["Friendly error screens active", "Empty states on all screens", "Retry widget functional"]),
        ("Day 48", "Performance Optimization",
         ["Profile all screens with Flutter DevTools to ensure 60fps",
          "Add const constructors to static layouts and components",
          "Optimize local list builders with lazy loading and pagination",
          "Lazy load cached images using cached_network_image package",
          "Reduce widget rebuilds by isolating Riverpod state select parameters"],
         ["60fps on all screens", "Paginated list builders active", "Images cached locally"]),
        ("Day 49", "Unit Testing — Mock Repositories",
         ["Mock repository interfaces using mocktail package",
          "Test MockAppointmentRepository: fetch, insert, update, cancel",
          "Test AuthNotifier: sign in, sign up, sign out state transitions",
          "Test MockProfileRepository: fetch profile, update profile",
          "Test BookingNotifier state machine (step 1 to 4 confirm flows)"],
         ["Repository tests passing", "AuthNotifier tests passing", "80%+ business logic coverage"]),
        ("Day 50", "Widget & Golden Testing",
         ["Write widget tests for custom design widgets (ZButton, ZCard, ZInput)",
          "Test booking flow screens using mocked Riverpod providers",
          "Test login form validations (empty, invalid email, short password)",
          "Perform golden tests for visual layout regression monitoring",
          "Test provider profile card rendering with mock datasets"],
         ["Widget tests complete", "Golden tests passing", "Validation tests active"]),
        ("Day 51", "Patrol E2E Testing",
         ["Write Patrol E2E test journey: signup, complete profile, book, pay",
          "Write Patrol E2E test: login, view bookings, cancel booking",
          "Run Patrol integration tests on physical iOS and Android devices",
          "Add E2E tests to GitHub Actions CI workflow runs"],
         ["3 E2E test journeys passing", "Real device tests completed", "CI running Patrol"]),
        ("Day 52", "Internal Beta Distribution",
         ["Build release IPAs (iOS) and App Bundles (AAB for Android)",
          "Upload to TestFlight and invite internal team (10 testers)",
          "Upload to Google Play Console Internal Testing track",
          "Configure Sentry SDK client-side crash reporting in main",
          "Collect structured feedback from internal team using form"],
         ["TestFlight build live", "Play Console internal track live", "Feedback form distributed"]),
        ("Day 53", "App Assets Preparation & Optimization",
         ["Audit app assets bundle size to minimize final application payload",
          "Optimize all PNG/JPEG/Lottie assets using compression tools",
          "Compress font files and remove unused font variations",
          "Clean up unused package dependencies in pubspec.yaml"],
         ["Assets optimized", "Package bundle size minimized", "Cleaner build output ready"]),
        ("Day 54", "External Beta — 50 Users",
         ["Expand TestFlight invite list to 50 external beta testers",
          "Expand Play Console closed track to 50 external beta testers",
          "Monitor Sentry crash logs for uncaught client exceptions",
          "Triage and resolve high-priority beta user bug reports",
          "Test performance on older physical target devices"],
         ["50 external testers active", "Sentry telemetry monitored", "Beta crash reports resolved"]),
        ("Day 55", "Local PHI Scrubbing & Log Security",
         ["Audit application logs to ensure no protected health info is outputted",
          "Configure local log rotation and diagnostic log file size limits",
          "Verify secure key storage configurations in Vault/Keyring platforms"],
         ["PHI data scrubbing active", "Local log security verified", "Encrypted log export ready"]),
        ("Day 56", "Offline Outbox Queue & API Webhook Failover",
         ["Build offline sync manager for outbox sync queue",
          "Write local db logic monitoring retry limits and sync statuses",
          "Verify failover handling under simulated network outages"],
         ["Outbox sync queue verified", "API failure logs active", "Failover scenarios tested"]),
        ("Day 57", "Hive Database Integrity Auditing",
         ["Perform database integrity checks on local Hive files during boot",
          "Audit local cache synchronization logs for structural discrepancies",
          "Design auto-repair logic resolving corrupted local cache files"],
         ["Hive integrity check active", "Cache database check completed", "Auto-repair routines verified"]),
        ("Day 58", "Buffer Day — Beta Bug Triaging & Crash Fixes",
         ["Review Sentry crash reports from active TestFlight and Play Console closed track",
          "Triage and fix critical visual bugs or state exceptions in booking flow",
          "Hotfix high-impact exception triggers in authentication and calendar widgets",
          "Build and deploy update builds to Apple TestFlight and Google Play Console"],
         ["All reported blocking exceptions resolved", "Updated beta builds successfully deployed"]),
        ("Day 59", "Buffer Day — Local Security Verification",
         ["Verify that all local database files use secure encryption keys",
          "Confirm that no patient information is outputted to device console logs",
          "Verify the end-to-end data export and delete GDPR flows",
          "Review App Store and Google Play data safety declaration questionnaires"],
         ["Security configuration verified", "Local logs audited and locked", "Data safety checklists ready"]),
        ("Day 60", "Compliance & Privacy Review",
         ["Ensure all local database encryption is verified",
          "Ensure no PHI is outputted to client-side logs or crash logs",
          "Add in-app privacy policy and terms of service screens",
          "Implement client-side GDPR data export utility",
          "Implement client-side GDPR data deletion utility"],
         ["Encryption verified", "Privacy policy screens in app", "GDPR data flows functional"]),
        ("Day 61", "Bug Fix Sprint",
         ["Triage all user feedback and bug reports from external beta",
          "Fix all crash exceptions on older devices",
          "UI fixes for small screen layout responsiveness",
          "Fix local notification scheduler triggers",
          "Upload final release candidate build to store consoles"],
         ["All crashes fixed", "UI layout bugs resolved", "Final release build ready"]),
        ("Day 62", "Release Candidate & Sign-Off",
         ["Full regression testing of all frontend client flows",
          "Tag release candidate in Git repository: v1.0.0-rc1",
          "Stakeholder sign-off meeting",
          "Freeze features - no new features until v1.1",
          "Prepare App Store and Google Play metadata, screenshots, and descriptions"],
         ["Release candidate tagged", "Stakeholder sign-off complete", "Store metadata ready"]),
    ]

    p5_data = [
        ("Day 63", "App Store Assets & Metadata",
         ["Design screenshots for all required sizes (6.7\", 6.1\", 5.5\", iPad)",
          "Write App Store description with mental health and therapy keywords",
          "Create 30-second app preview demonstration video",
          "Finalize 1024x1024px App Store icon",
          "Live privacy policy and terms of service URLs on website"],
         ["Screenshots ready", "App Store metadata copy finalized", "Privacy link live"]),
        ("Day 64", "iOS App Store Submission",
         ["Configure Xcode signing and distribution credentials",
          "Build release IPA using command: flutter build ipa --release",
          "Upload IPA to App Store Connect using Fastlane tools",
          "Complete App Store Connect review questionnaires (encryption, data safety)",
          "Submit for Apple review"],
         ["IPA uploaded to App Store Connect", "Connect questionnaires complete", "Submitted to Apple for review"]),
        ("Day 65", "Google Play Submission",
         ["Build Android App Bundle: flutter build appbundle --release",
          "Sign with release keystore credentials",
          "Upload App Bundle (AAB) to Google Play Console Production track",
          "Complete Google Play Store listing information",
          "Publish build for Google Play review"],
         ["AAB uploaded to Play Console", "Play Store listing complete", "Live on Google Play"]),
        ("Day 66", "Review Delay Buffer — App Store Feedbacks",
         ["Monitor App Store Connect review progress and respond immediately to reviewer inquiries",
          "Address potential app classification queries regarding medical category",
          "Prepare alternative metadata descriptions if requested by reviewers",
          "Configure pricing tiers and test-user credentials in App Store Connect"],
         ["Apple reviewer feedback handled", "iOS store requirements met"]),
        ("Day 67", "Review Delay Buffer — Google Play Console Closed Testing",
         ["Verify that closed-testing compliance tracks remain active with required testers",
          "Monitor Google Play dashboard for console alerts or review delays",
          "Review pre-launch reports generated by the Play Console test lab",
          "Prepare fallback builds in case of build-rejection notifications"],
         ["Google pre-launch alerts resolved", "Android store requirements met"]),
        ("Day 68", "Sentry & Mixpanel Setup",
         ["Enable Sentry crash reporting configuration in the production app",
          "Configure Sentry alerts for new crash types and anomalies",
          "Set up Mixpanel events: app_open, booking_started, payment_success",
          "Create Mixpanel funnel: Provider View to Book to Pay to Confirm"],
         ["Sentry alerts configured", "Mixpanel funnels active", "Production monitoring live"]),
        ("Day 69", "App Store Optimization & Marketing",
         ["Research top search keywords for mental health apps",
          "Optimize App Store title and subtitle with target keywords",
          "Set up Apple Search Ads campaign",
          "Set up Google UAC campaign for Android installs",
          "Announce launch on LinkedIn, Instagram, and health tech communities"],
         ["Listings optimized", "Marketing campaigns active", "Launch announced"]),
        ("Day 70", "Launch Day Monitoring",
         ["Monitor Sentry dashboard for any crash reports from production users",
          "Watch for new reviews on App Store and Google Play Console",
          "Monitor booking conversion rates in Mixpanel funnels",
          "Triage and prepare hotfixes if any P0 crashes occur"],
         ["Launch stable", "Sentry monitored", "Conversion rates tracked"]),
        ("Day 71", "Post-Launch Data Analysis",
         ["Analyze first 7 days: installs, DAU, booking conversion, drop-offs",
          "Query local Hive stats database to find popular provider categories",
          "Identify top 3 friction points in user experience funnels",
          "Survey first 50 active users for qualitative feedback",
          "Document findings in v1.1 planning documentation"],
         ["User behavior report written", "Top friction points identified", "v1.1 planning doc started"]),
        ("Day 72", "Retrospective & v1.1 Roadmap",
         ["Full project retrospective with team (what went well, what to improve)",
          "Publish v1.1 public roadmap overview",
          "Set 90-day OKRs: target bookings, DAU, retention rates",
          "Start v1.1 backlog: telehealth group sessions, Apple Health integration",
          "Celebrate launch - 75 days complete!"],
         ["Retrospective complete", "v1.1 roadmap published", "OKRs set for next quarter"]),
        ("Day 73", "Post-Launch Buffer — Performance Tuning",
         ["Monitor application frame rates and page rendering performance on older devices",
          "Analyze Sentry transaction metrics to find slow rendering states",
          "Verify Sentry/Mixpanel analytics webhook performance in production"],
         ["Telemetry analyzed", "Performance tuning complete"]),
        ("Day 74", "Post-Launch Buffer — Production Feedback",
         ["Monitor customer support channels for login or booking layout complaints",
          "Draft and verify hotfixes for production styling bugs or local storage issues",
          "Update App Store and Google Play Console support resources and links"],
         ["Initial user issues triaged", "Minor updates rolled out"]),
        ("Day 75", "Buffer Day — Handoff & Backlog Grooming",
         ["Conduct a final engineering retrospective and update code documentation in the repository",
          "Clean up development and staging local configuration settings",
          "Prioritize features in the backlog for the upcoming v1.1 version cycle",
          "Prepare system administration handover documentation for the client team"],
         ["Engineering retro summary written", "Roadmap backlog prioritized", "Handover documentation ready"]),
    ]

    # Helper function to append daily items in docx
    def add_phase_days(title_text, days_data):
        doc.add_heading(title_text, level=2)
        for day, heading, tasks, deliverables in days_data:
            p_day = doc.add_paragraph()
            p_day.add_run(f"{day}: {heading}").bold = True
            p_day.runs[0].font.size = Pt(12)
            
            p_tasks = doc.add_paragraph()
            p_tasks.add_run("Tasks:").italic = True
            for t in tasks:
                doc.add_paragraph(t, style='List Bullet')
                
            p_deliv = doc.add_paragraph()
            p_deliv.add_run("Deliverables:").italic = True
            for d in deliverables:
                doc.add_paragraph(d, style='List Bullet')
                
            doc.add_paragraph() # Spacer

    add_phase_days("Phase 1 — Foundation & Setup (Days 1–9)", p1_data)
    doc.add_page_break()
    add_phase_days("Phase 2 — Core Features  (Days 10–26)", p2_data)
    doc.add_page_break()
    add_phase_days("Phase 3 — Advanced Features & Local Caching (Days 27–45)", p3_data)
    doc.add_page_break()
    add_phase_days("Phase 4 — Polish & Testing (Days 46–62)", p4_data)
    doc.add_page_break()
    add_phase_days("Phase 5 — Launch & Deploy (Days 63–75)", p5_data)
    doc.add_page_break()

    # 4. DAILY TRACKER TABLE
    doc.add_heading("Daily Deliverables Tracker", level=1)
    doc.add_paragraph("Check off each day as it is completed.")
    
    tracker_table = doc.add_table(rows=1, cols=4)
    tracker_table.style = 'Table Grid'
    tracker_hdr = tracker_table.rows[0].cells
    tracker_hdr[0].text = "Day"
    tracker_hdr[1].text = "Ph"
    tracker_hdr[2].text = "Deliverable Description"
    tracker_hdr[3].text = "Done"
    for cell in tracker_hdr:
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_background(cell, "E8F7F0")

    # Reconstruct tracking deliverables
    all_days = p1_data + p2_data + p3_data + p4_data + p5_data
    phase_mapping = {}
    for d_num in range(1, 10): phase_mapping[d_num] = "1"
    for d_num in range(10, 27): phase_mapping[d_num] = "2"
    for d_num in range(27, 46): phase_mapping[d_num] = "3"
    for d_num in range(46, 63): phase_mapping[d_num] = "4"
    for d_num in range(63, 76): phase_mapping[d_num] = "5"

    for idx, (day_str, heading, _, deliverables) in enumerate(all_days):
        day_num = idx + 1
        ph_str = phase_mapping.get(day_num, "1")
        primary_deliverable = deliverables[0] if deliverables else heading
        
        row_cells = tracker_table.add_row().cells
        row_cells[0].text = str(day_num)
        row_cells[1].text = ph_str
        row_cells[2].text = primary_deliverable
        row_cells[3].text = "[  ]"
        
    doc.add_page_break()

    # 5. COSTS TABLE
    doc.add_heading("Estimated Project Costs", level=1)
    doc.add_paragraph(
        "Direct monetary costs required to publish, deploy, host, "
        "and maintain the Zenup Health mobile application."
    )
    
    costs_table = doc.add_table(rows=1, cols=4)
    costs_table.style = 'Table Grid'
    costs_hdr = costs_table.rows[0].cells
    costs_hdr[0].text = "Item / Service"
    costs_hdr[1].text = "Billing Frequency"
    costs_hdr[2].text = "Estimated Cost"
    costs_hdr[3].text = "Purpose / Notes"
    for cell in costs_hdr:
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_background(cell, "E8F7F0")

    costs_raw = [
        ['Google Play Console', 'One-time', '$25 USD', 'Required developer account for Android publishing.'],
        ['Apple Developer Program', 'Annual', '$99 USD / year', 'Required developer account for iOS publishing.'],
        ['MediaSFU Meetings', 'Monthly (Usage)', '$162 USD / month', 'Client video calling sessions usage (10% discount).'],
        ['Razorpay Gateway', 'Transaction-based', '2.0% per transaction (+ GST)', 'Domestic payment processing fee for client transactions.'],
        ['Sentry Error Tracker', 'Free Tier', '$0 USD', 'Client-side crash and error tracking.'],
    ]
    for row in costs_raw:
        row_cells = costs_table.add_row().cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1]
        row_cells[2].text = row[2]
        row_cells[3].text = row[3]

    doc.add_page_break()

    # 6. RISK REGISTER
    doc.add_heading("Risk Register", level=1)
    
    risks_table = doc.add_table(rows=1, cols=4)
    risks_table.style = 'Table Grid'
    risks_hdr = risks_table.rows[0].cells
    risks_hdr[0].text = "Risk"
    risks_hdr[1].text = "Level"
    risks_hdr[2].text = "Cause"
    risks_hdr[3].text = "Mitigation"
    for cell in risks_hdr:
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_background(cell, "E8F7F0")

    risks = [
        ["Apple Review Rejection", "High", "Missing privacy declarations for health data", "Add NSHealthShareUsageDescription, privacy policy URL, and data use disclosures before Day 43"],
        ["Local PHI Storage Security", "High", "Local patient data stored insecurely on device", "Encrypt Hive database boxes using keys stored in keychain/keystore using secure storage"],
        ["Razorpay SDK Timeout", "Medium", "Client-side payment gateway fails to connect or times out", "Implement robust state retry mechanisms and check connectivity state dynamically"],
        ["MediaSFU Room Expiry", "Medium", "Call drops if session credentials expire", "Implement client-side token refresh and auto-reconnect flow in video controller"],
        ["In-App Performance", "Medium", "High memory usage from video/charts leads to app crashes on older hardware", "Profile memory with DevTools, dispose controllers, and paginate local list builders"],
    ]
    for row in risks:
        row_cells = risks_table.add_row().cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1]
        row_cells[2].text = row[2]
        row_cells[3].text = row[3]
        if row[1] == "High":
            set_cell_background(row_cells[1], "FADBD8") # Light Red
        elif row[1] == "Medium":
            set_cell_background(row_cells[1], "FCF3CF") # Light Yellow
        else:
            set_cell_background(row_cells[1], "D5F5E3") # Light Green

    doc.add_paragraph()

    # 7. LAUNCH CHECKLIST
    doc.add_heading("Launch Checklist", level=1)
    
    cl = [
        ("App Store (iOS)", [
            "App icon 1024x1024px — no alpha, no rounded corners (Apple adds them)",
            "Screenshots for 6.7\", 6.1\", 5.5\", 12.9\" iPad",
            "Age rating 17+ selected (health / medical content)",
            "Privacy policy and terms of service URLs in App Store Connect",
            "NSHealthShareUsageDescription in Info.plist if using HealthKit",
            "Apple Sign-In implemented (required when any other OAuth is present)",
            "TestFlight tested with 50+ users before production submission",
        ]),
        ("Google Play (Android)", [
            "Feature graphic 1024x500px",
            "Screenshots for phone and 7\" + 10\" tablet",
            "Content rating questionnaire complete (Health category)",
            "Target SDK 34+ (Android 14 requirement)",
            "App Bundle (AAB) uploaded — NOT APK",
            "Data safety section filled (data collected, shared, encrypted)",
            "Closed testing with 20+ testers before production",
        ]),
        ("Flutter App Hardening", [
            "ProGuard rules configured for Android release build",
            "iOS bitcode disabled (deprecated — ensure Xcode setting correct)",
            "Client API keys stored securely in native config or env files",
            "Certificate pinning implemented for MediaSFU and Razorpay API calls",
            "Biometric auth gate on sensitive screens",
            "Sentry or Crashlytics configured for production crash tracking",
            "App version and build number incremented (pubspec.yaml)",
        ]),
    ]
    for section, items in cl:
        doc.add_heading(section, level=3)
        for item in items:
            doc.add_paragraph(f"[  ] {item}")
            
    out_path = os.path.join(os.path.dirname(__file__), "Zenup_Health_75Day_Plan_v3.docx")
    doc.save(out_path)
    print(f"Word document saved successfully to: {out_path}")

if __name__ == "__main__":
    generate_document()
